from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Union
from io import BytesIO
from docx.shared import Pt, RGBColor


async def generate_document(json_data: str, template_file: Union[BytesIO, str], output_file: str):
    """
    Generates a DOCX file based on a template and JSON data.

    Args:
        json_data (dict): JSON data as a dictionary.
        template_file (BytesIO): In-memory DOCX template file.
        output_file (str): Path to save the generated document.

    Returns:
        None
    """
    # Load the template
    template_doc = Document(template_file)

    # Extract month and sections
    month = json_data.get("document", {}).get("month", "")
    sections = json_data.get("document", {}).get("sections", [])

    # First pass: Replace {date} placeholder with month value
    for paragraph in template_doc.paragraphs:
        if "{{date}}" in paragraph.text:
            for run in paragraph.runs:
                if "{{date}}" in run.text:
                    run.text = run.text.replace("{{date}}", month)

    # Process sections as before
    for section in sections:
        title = section["title"]
        bullets = section["bullets"]

        # Locate placeholder for the current section
        for paragraph in template_doc.paragraphs:
            if f"{{{{{title}}}}}" in paragraph.text:
                paragraph.text = ""  # Clear the placeholder text

                # Insert bullets for the section
                for bullet in bullets:
                    bullet_paragraph = paragraph.insert_paragraph_before()
                    bullet_paragraph.style = "List Bullet"

                    run = None  # Initialize `run` to avoid referencing before assignment
                    # Add the hyperlink or styled text
                    if "link" in bullet and bullet["link"]:
                        add_hyperlink(
                            paragraph=bullet_paragraph, 
                            text=bullet["text"], 
                            url=bullet["link"], 
                            styles=bullet.get("styles", [])
                        )
                    else:
                        run = bullet_paragraph.add_run(bullet["text"])
                        _apply_styles(run, bullet.get("styles", []))

                    # Add nested content below the bullet if it exists
                    if "content" in bullet and bullet["content"].strip():
                        content_paragraph = paragraph.insert_paragraph_before()
                        content_paragraph.style = "Normal"  # Content below bullets shouldn't have a bullet style
                        
                        # Format date and content if date exists
                        date_text = ""
                        if "date" in bullet:
                            date_text = _format_date(bullet["date"]) + " "
                        
                        content_run = content_paragraph.add_run(f"{date_text}[…] {bullet['content']}")
                        _apply_blue_style(content_run)

                        # Align content dynamically with the bullet
                        bullet_indent = bullet_paragraph.paragraph_format.left_indent or Pt(18)
                        content_paragraph.paragraph_format.left_indent = bullet_indent  # Match bullet's indent
                        content_paragraph.paragraph_format.first_line_indent = Pt(0)  # No extra indentation

                        # Add space after content only when content exists
                        content_paragraph.paragraph_format.space_after = Pt(12)

                    # Apply blue style only to the `run` for plain text or styled bullets
                    if run:
                        _apply_blue_style(run)
                break

    # Final pass: Clear any remaining placeholders
    for paragraph in template_doc.paragraphs:
        text = paragraph.text
        if "{{" in text and "}}" in text:
            # Find all placeholder patterns and replace them with empty string
            while "{{" in text and "}}" in text:
                start = text.find("{{")
                end = text.find("}}") + 2
                text = text[:start] + text[end:]
            paragraph.text = text

    # Save the updated document
    template_doc.save(output_file)


def add_hyperlink(paragraph, text, url, styles=None):
    """
    Add a hyperlink to a paragraph with optional styles (bold, italic, underline).

    Args:
        paragraph: The paragraph to add the hyperlink to.
        text: The display text for the hyperlink.
        url: The URL for the hyperlink.
        styles: List of styles to apply (e.g., ["bold", "italic", "underline"]).

    Returns:
        None
    """
    if not url:
        raise ValueError("URL for hyperlink cannot be None")

    if styles is None:
        styles = []

    # Create the hyperlink XML element
    hyperlink = OxmlElement("w:hyperlink")
    rId = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink.set(qn("r:id"), rId)

    # Create a run element for the link text
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")  # Run properties

    # Apply standard hyperlink style
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # Set the font color to blue
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0033a0")  # Blue
    rPr.append(color)

    # Apply additional styles (bold, italic, underline)
    if "bold" in styles:
        bold = OxmlElement("w:b")
        rPr.append(bold)
    if "italic" in styles:
        italic = OxmlElement("w:i")
        rPr.append(italic)
    if "underline" in styles:
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)

    run.append(rPr)

    # Add the text element
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)

    # Append the hyperlink to the paragraph
    paragraph._element.append(hyperlink)


def _apply_styles(run, styles):
    """
    Apply styles (bold, italic, underline) to a text run.

    Args:
        run: The text run to style.
        styles: List of styles to apply (e.g., ["bold", "italic", "underline"]).

    Returns:
        None
    """
    if "bold" in styles:
        run.bold = True
    if "italic" in styles:
        run.italic = True
    if "underline" in styles:
        run.underline = True


def _apply_blue_style(run):
    """
    Apply blue styling to a text run.

    Args:
        run: The text run to style.

    Returns:
        None
    """
    run.font.color.rgb = RGBColor(0, 51, 160)  # Blue
    run.font.size = Pt(11)  # Match standard font size


def _apply_normal_style(run):
    """
    Apply normal text styling (for nested content).

    Args:
        run: The text run to style.

    Returns:
        None
    """
    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
    run.font.size = Pt(11)  # Match standard font size


def _format_date(date_str):
    """
    Convert date string from 'Month DD, YYYY' to '[DD Month]' format
    """
    from datetime import datetime
    try:
        date_obj = datetime.strptime(date_str, '%B %d, %Y')
        return f"[{date_obj.day} {date_obj.strftime('%B')}]"
    except:
        return ""